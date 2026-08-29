"""
Reading a circular out of the file it arrived in. Part G.

The rule this module works under
---------------------------------
**An extractor that cannot read a document says so.** It never returns empty
text and lets the caller conclude the circular imposes no obligations. That
distinction is the whole design: a circular with no rules and a circular
nobody could read look identical in a corpus, and only one of them is a
finding.

So every extractor declares whether it is available on this deployment, and
every result carries the reason it produced what it produced. A PDF whose
pages are scans comes back `needs_ocr`; a format with no extractor comes back
`unavailable`; a document that genuinely says nothing comes back extracted and
empty, which is a different and much rarer thing.

OCR
---
An OCR engine is a system binary, not a Python package, and a bank's image may
or may not carry one. `ocr_available()` reports the truth, and a page with no
text is marked NEEDS_OCR rather than silently dropped. Nothing here shells out
to a binary it has not confirmed exists.

What is deliberately NOT done
------------------------------
No model reads the document. Extraction is deterministic: page text, section
numbering, tables, and pattern-matched candidate rules. The obligations a
model would find are exactly the ones nobody could check, and Part G's whole
value is that a citation resolves to a page in a file with a known hash.

The candidate rules are candidates
-----------------------------------
A sentence containing "shall" is *probably* an obligation. `_obligation` says
so with a confidence and a reason, and an SME decides. Extraction that
promoted its own guesses to approved rules would be the regulatory equivalent
of an auto-approved teaching case.
"""

from __future__ import annotations

import csv
import io
import logging
import re
import shutil
from dataclasses import dataclass, field
from typing import Any

from backend.regulatory import schema as rs

logger = logging.getLogger(__name__)

EXTRACT_VERSION = "1.0.0"


@dataclass
class Extraction:
    """What came out of one file, and how much of it is trustworthy."""

    pages: list[rs.Page] = field(default_factory=list)
    sections: list[rs.Section] = field(default_factory=list)
    rules: list[rs.Rule] = field(default_factory=list)
    status: str = rs.EXTRACTED
    #: Why this status. Shown to the person who uploaded it.
    because: str = ""
    #: Which extractor ran, so a re-extraction after a library upgrade can be
    #: told apart from the first attempt.
    extractor: str = ""
    version: str = EXTRACT_VERSION

    @property
    def text(self) -> str:
        return "\n".join(p.text for p in self.pages if p.text)

    @property
    def pages_needing_ocr(self) -> list[int]:
        return [p.number for p in self.pages if p.needs_ocr]

    def to_dict(self) -> dict[str, Any]:
        return {"pages": [p.to_dict() for p in self.pages],
                "sections": [s.to_dict() for s in self.sections],
                "rules": [r.to_dict() for r in self.rules],
                "status": self.status, "because": self.because,
                "extractor": self.extractor, "version": self.version,
                "page_count": len(self.pages),
                "characters": len(self.text),
                "pages_needing_ocr": self.pages_needing_ocr}


# ---------------------------------------------------------------------------
# Availability
# ---------------------------------------------------------------------------


def pdf_available() -> bool:
    try:
        import pypdf  # noqa: F401
    except Exception:  # noqa: BLE001 - absence is the answer, not a failure
        return False
    return True


def docx_available() -> bool:
    try:
        import docx  # noqa: F401
    except Exception:  # noqa: BLE001
        return False
    return True


def xlsx_available() -> bool:
    try:
        import openpyxl  # noqa: F401
    except Exception:  # noqa: BLE001
        return False
    return True


def ocr_available() -> bool:
    """Whether this deployment can read a scanned page.

    Two things have to be true: the Python binding and the engine binary. A
    deployment with `pytesseract` installed and no `tesseract` on PATH raises
    at the first page, which is worse than knowing in advance.
    """
    try:
        import pytesseract  # noqa: F401
    except Exception:  # noqa: BLE001
        return False
    return shutil.which("tesseract") is not None


def availability() -> dict[str, Any]:
    """What this deployment can and cannot read, for the upload screen."""
    formats = {
        rs.PDF: pdf_available(),
        rs.DOCX: docx_available(),
        rs.XLSX: xlsx_available(),
        rs.CSV: True,
        rs.TXT: True,
        rs.HTML: True,
    }
    return {
        "formats": formats,
        "unavailable": sorted(f for f, ok in formats.items() if not ok),
        "ocr": ocr_available(),
        "ocr_note": (
            "Scanned pages can be read on this deployment."
            if ocr_available() else
            "No OCR engine is configured, so a scanned page is recorded as "
            "NEEDS_OCR rather than read. Its obligations are not in the "
            "corpus and nothing pretends otherwise."),
        "version": EXTRACT_VERSION,
    }


# ---------------------------------------------------------------------------
# Per-format extraction
# ---------------------------------------------------------------------------


def _pdf(payload: bytes) -> Extraction:
    if not pdf_available():
        return Extraction(
            status=rs.EXTRACTION_UNAVAILABLE, extractor="pdf",
            because="No PDF reader is installed on this deployment.")
    import pypdf

    try:
        reader = pypdf.PdfReader(io.BytesIO(payload))
    except Exception as e:  # noqa: BLE001
        return Extraction(
            status=rs.EXTRACTION_UNAVAILABLE, extractor="pdf",
            because=f"The PDF could not be opened: {e}")

    pages: list[rs.Page] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as e:  # noqa: BLE001 - one bad page, not the file
            logger.warning("Page %s could not be read: %s", index, e)
            text = ""
        pages.append(rs.Page(number=index, text=text.strip(),
                             needs_ocr=not text.strip()))

    scanned = [p for p in pages if p.needs_ocr]
    if scanned and len(scanned) == len(pages):
        return Extraction(
            pages=pages, status=rs.NEEDS_OCR, extractor="pdf",
            because=(f"All {len(pages)} pages carry no extractable text. "
                     "They are images." + ("" if ocr_available() else
                     " No OCR engine is configured on this deployment.")))
    because = f"{len(pages)} page(s) read."
    if scanned:
        because += (f" {len(scanned)} of them carry no extractable text and "
                    "are marked NEEDS_OCR.")
    return Extraction(pages=pages, status=rs.EXTRACTED, extractor="pdf",
                      because=because)


def _docx(payload: bytes) -> Extraction:
    if not docx_available():
        return Extraction(
            status=rs.EXTRACTION_UNAVAILABLE, extractor="docx",
            because="No DOCX reader is installed on this deployment.")
    import docx

    try:
        document = docx.Document(io.BytesIO(payload))
    except Exception as e:  # noqa: BLE001
        return Extraction(
            status=rs.EXTRACTION_UNAVAILABLE, extractor="docx",
            because=f"The DOCX could not be opened: {e}")

    body = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    tables = [[[cell.text.strip() for cell in row.cells] for row in t.rows]
              for t in document.tables]
    # A Word file has no pages until it is laid out. One page, honestly
    # numbered 1, rather than a page count invented from a character estimate.
    page = rs.Page(number=1, text=body, tables=tables,
                   needs_ocr=not body.strip() and not tables)
    return Extraction(
        pages=[page], status=rs.EXTRACTED, extractor="docx",
        because=(f"{len(document.paragraphs)} paragraph(s) and "
                 f"{len(tables)} table(s) read. A DOCX has no pages until it "
                 "is laid out, so citations carry the section rather than a "
                 "page number."))


def _xlsx(payload: bytes) -> Extraction:
    if not xlsx_available():
        return Extraction(
            status=rs.EXTRACTION_UNAVAILABLE, extractor="xlsx",
            because="No XLSX reader is installed on this deployment.")
    import openpyxl

    try:
        book = openpyxl.load_workbook(io.BytesIO(payload), data_only=True,
                                      read_only=True)
    except Exception as e:  # noqa: BLE001
        return Extraction(
            status=rs.EXTRACTION_UNAVAILABLE, extractor="xlsx",
            because=f"The workbook could not be opened: {e}")

    pages: list[rs.Page] = []
    for index, name in enumerate(book.sheetnames, start=1):
        sheet = book[name]
        rows = [[("" if cell is None else str(cell)) for cell in row]
                for row in sheet.iter_rows(values_only=True)]
        rows = [r for r in rows if any(c.strip() for c in r)]
        text = "\n".join("\t".join(r) for r in rows)
        pages.append(rs.Page(number=index, text=f"{name}\n{text}",
                             tables=[rows] if rows else []))
    book.close()
    return Extraction(pages=pages, status=rs.EXTRACTED, extractor="xlsx",
                      because=f"{len(pages)} sheet(s) read as tables.")


def _csv(payload: bytes) -> Extraction:
    text = payload.decode("utf-8", errors="replace")
    rows = [r for r in csv.reader(io.StringIO(text)) if any(c.strip()
                                                            for c in r)]
    page = rs.Page(number=1, text="\n".join("\t".join(r) for r in rows),
                   tables=[rows] if rows else [])
    return Extraction(pages=[page], status=rs.EXTRACTED, extractor="csv",
                      because=f"{len(rows)} row(s) read.")


def _txt(payload: bytes) -> Extraction:
    text = payload.decode("utf-8", errors="replace")
    return Extraction(
        pages=[rs.Page(number=1, text=text.strip(),
                       needs_ocr=not text.strip())],
        status=rs.EXTRACTED, extractor="txt",
        because=f"{len(text)} character(s) read.")


#: Tags whose CONTENT is not document text. A circular's obligations are not
#: in its stylesheet, and leaving them in produced "candidate obligations"
#: made of CSS.
_DROPPED = re.compile(r"<(script|style)[^>]*>.*?</\1>",
                      re.IGNORECASE | re.DOTALL)
_TAG = re.compile(r"<[^>]+>")
_ENTITY = {"&nbsp;": " ", "&amp;": "&", "&lt;": "<", "&gt;": ">",
           "&quot;": '"', "&#39;": "'"}


def _html(payload: bytes) -> Extraction:
    """Tags stripped deterministically rather than parsed by a library.

    A regulator's HTML is a document, not an application, and the two things
    that matter — the running text and the tables — survive tag removal. This
    keeps a parser dependency out of the image for a format that does not need
    one.
    """
    raw = payload.decode("utf-8", errors="replace")
    raw = _DROPPED.sub(" ", raw)
    raw = re.sub(r"</(p|div|tr|h[1-6]|li)>", "\n", raw, flags=re.IGNORECASE)
    raw = re.sub(r"</t[dh]>", "\t", raw, flags=re.IGNORECASE)
    text = _TAG.sub("", raw)
    for entity, character in _ENTITY.items():
        text = text.replace(entity, character)
    lines = [ln.strip() for ln in text.splitlines()]
    body = "\n".join(ln for ln in lines if ln)
    return Extraction(
        pages=[rs.Page(number=1, text=body, needs_ocr=not body)],
        status=rs.EXTRACTED, extractor="html",
        because=f"{len(body)} character(s) read after tag removal.")


_EXTRACTORS = {
    rs.PDF: _pdf, rs.DOCX: _docx, rs.XLSX: _xlsx,
    rs.CSV: _csv, rs.TXT: _txt, rs.HTML: _html,
}


# ---------------------------------------------------------------------------
# Sections
# ---------------------------------------------------------------------------

#: "4.", "4.2", "4.2.1", "Article 7", "Part III", "Section 12". A citation
#: points at one of these, so the numbering has to be recognised rather than
#: guessed from indentation.
_NUMBERED = re.compile(
    r"^\s*(?:(?:Article|Section|Part|Clause|Paragraph|Chapter)\s+)?"
    r"((?:\d+|[IVXLC]+)(?:\.\d+)*)\.?\s+(\S.*)$", re.IGNORECASE)

#: How long a line can be and still be a heading. A 300-character "heading" is
#: a paragraph whose first words happened to look like a number.
_HEADING_LIMIT = 160


#: How long the remainder of a numbered line can be and still be read as a
#: heading rather than as the provision itself.
_HEADING_WORDS = 12


def _is_heading(remainder: str) -> bool:
    """Whether "4. Scope" is a heading and "4.1 A bank shall..." is not.

    This distinction was worth getting wrong once to learn. Treating every
    numbered line as a heading put the provision INTO the heading and left the
    section text empty — so a circular with five obligations produced five
    sections and zero rules, and the corpus looked structured and said
    nothing.

    A regulator numbers headings AND provisions with the same notation. The
    difference is grammatical: a heading is a short noun phrase with no final
    stop; a provision is a sentence.
    """
    text = remainder.strip()
    if not text:
        return False
    if text.endswith((".", ";", ":")):
        return False
    return len(text.split()) <= _HEADING_WORDS


def sections_of(pages: list[rs.Page]) -> list[rs.Section]:
    """The numbered structure of the document, in order.

    Deterministic and conservative. A line beginning with recognised numbering
    opens a section; the rest of that line is the section's heading when it
    reads as one and the first line of its text when it does not. Everything
    up to the next numbered line belongs to it.
    """
    found: list[rs.Section] = []
    current: rs.Section | None = None
    body: list[str] = []

    def close() -> None:
        if current is not None:
            current.text = "\n".join(body).strip()
            found.append(current)

    for page in pages:
        for line in page.text.splitlines():
            match = _NUMBERED.match(line)
            if match and len(line) <= _HEADING_LIMIT:
                close()
                number = match.group(1)
                remainder = match.group(2).strip()
                heading = remainder if _is_heading(remainder) else ""
                body = [] if heading else [remainder]
                parent = number.rsplit(".", 1)[0] if "." in number else ""
                current = rs.Section(
                    section_id=f"s{len(found) + 1}", number=number,
                    heading=heading, page=page.number, parent=parent)
                continue
            if line.strip():
                body.append(line.strip())
    close()
    return found


# ---------------------------------------------------------------------------
# Candidate rules
# ---------------------------------------------------------------------------

#: Modal verbs a regulator uses to impose. "Should" is deliberately absent:
#: it is guidance in most rulebooks, and treating guidance as an obligation
#: makes the corpus over-claim in the direction nobody checks.
_MUST = re.compile(
    r"\b(shall not|must not|may not|shall|must|is required to|are required "
    r"to|is prohibited|are prohibited)\b", re.IGNORECASE)

_DEFINES = re.compile(
    r"\b(means|shall mean|is defined as|refers to|for the purposes of this)\b",
    re.IGNORECASE)

_EXCEPTS = re.compile(
    r"\b(except|unless|does not apply|shall not apply|other than|save for|"
    r"exempt)\b", re.IGNORECASE)

#: A number with a unit a treatment can turn on.
#:
#: The trailing word boundary is inside the alternation, not after it. Written
#: as `(%|per cent|...)\b` it silently matched nothing ending in `%`: `\b`
#: after a non-word character requires a word character next, and "1.5 % of"
#: has a space there. Every percentage threshold in every circular would have
#: been missed, and the corpus would have looked like it had simply found no
#: percentages.
_NUMBER = re.compile(
    r"(?<![\w.])(\d{1,3}(?:,\d{3})*(?:\.\d+)?|\d+(?:\.\d+)?)\s*"
    r"(%|(?:per cent|percent|basis points|bps|days|months|years|times|x)\b)",
    re.IGNORECASE)

_UNITS = {"per cent": "%", "percent": "%", "basis points": "bps",
          "bps": "bps", "days": "days", "months": "months", "years": "years",
          "times": "x", "x": "x", "%": "%"}

#: Sentence splitting that does not break on "4.2." or "SR 1.5 million".
_SENTENCE = re.compile(r"(?<=[.;:])\s+(?=[A-Z(])")

#: A sentence shorter than this is a fragment — a heading, a page number, a
#: table cell — and proposing it as an obligation gives a reviewer noise to
#: wade through instead of rules to judge.
_MIN_RULE = 40
_MAX_RULE = 1200


def _clean(sentence: str) -> str:
    return re.sub(r"\s+", " ", sentence).strip()


def rules_of(sections: list[rs.Section], pages: list[rs.Page],
             *, concepts: tuple[str, ...] = ()) -> list[rs.Rule]:
    """Candidate obligations, definitions, thresholds and exceptions.

    Every one carries its section, its page, why it was proposed, and a
    confidence. None of them is approved by producing it.

    A sentence may yield more than one rule — an obligation that also states a
    threshold is both, and splitting them would lose the obligation the
    threshold belongs to. So the threshold carries the same text.
    """
    out: list[rs.Rule] = []

    def emit(kind: str, text: str, section: rs.Section | None, page: int,
             *, because: str, confidence: float,
             value: float | None = None, unit: str = "") -> None:
        out.append(rs.Rule(
            rule_id=f"r{len(out) + 1}", kind=kind, text=text,
            section_id=getattr(section, "section_id", ""),
            section_number=getattr(section, "number", ""),
            page=page or getattr(section, "page", 0),
            value=value, unit=unit,
            concepts=[c for c in concepts if c.lower() in text.lower()],
            because=because, confidence=confidence)) 

    units = [(s, s.text, s.page) for s in sections]
    if not units:
        # A document with no recognised numbering is still a document. Its
        # rules cite the page instead of the section, which is worse and is
        # not nothing.
        units = [(None, p.text, p.number) for p in pages]

    for section, text, page in units:
        for raw in _SENTENCE.split(text or ""):
            sentence = _clean(raw)
            if not (_MIN_RULE <= len(sentence) <= _MAX_RULE):
                continue

            modal = _MUST.search(sentence)
            defines = _DEFINES.search(sentence)
            excepts = _EXCEPTS.search(sentence)
            numbers = list(_NUMBER.finditer(sentence))

            if excepts:
                emit(rs.EXCEPTION, sentence, section, page,
                     because=(f"the sentence carries "
                              f"{excepts.group(1).lower()!r}, which is how a "
                              "carve-out is written"),
                     confidence=0.6)
            elif modal:
                emit(rs.OBLIGATION, sentence, section, page,
                     because=(f"the sentence carries the modal "
                              f"{modal.group(1).lower()!r}"),
                     confidence=0.75 if modal.group(1).lower() in
                     ("shall", "must", "shall not", "must not") else 0.6)
            elif defines:
                emit(rs.DEFINITION, sentence, section, page,
                     because=(f"the sentence carries "
                              f"{defines.group(1).lower()!r}, which is how a "
                              "term is defined"),
                     confidence=0.65)

            for number in numbers:
                raw_value = number.group(1).replace(",", "")
                try:
                    value = float(raw_value)
                except ValueError:  # pragma: no cover - the regex guarantees
                    continue
                unit = _UNITS.get(number.group(2).lower(), number.group(2))
                emit(rs.THRESHOLD, sentence, section, page,
                     because=(f"the sentence states {number.group(0)}, which "
                              "a treatment can turn on"),
                     confidence=0.7, value=value, unit=unit)
    return out


# ---------------------------------------------------------------------------
# The whole job
# ---------------------------------------------------------------------------


def extract(payload: bytes, file_format: str, *,
            concepts: tuple[str, ...] = ()) -> Extraction:
    """Read one original. Never raises for a document it cannot read."""
    reader = _EXTRACTORS.get(file_format)
    if reader is None:
        return Extraction(
            status=rs.EXTRACTION_UNAVAILABLE, extractor="",
            because=(f"{file_format!r} is not a format CreditProbe reads. "
                     "Supported: " + ", ".join(rs.FORMATS) + "."))
    try:
        found = reader(payload)
    except Exception as e:  # noqa: BLE001 - an unreadable file is a status
        logger.warning("Extraction failed for a %s: %s", file_format, e)
        return Extraction(status=rs.EXTRACTION_UNAVAILABLE,
                          extractor=file_format.lower(),
                          because=f"The file could not be read: {e}")

    if found.status in (rs.EXTRACTION_UNAVAILABLE, rs.NEEDS_OCR):
        return found

    found.sections = sections_of(found.pages)
    found.rules = rules_of(found.sections, found.pages, concepts=concepts)
    found.because += (f" {len(found.sections)} section(s) and "
                      f"{len(found.rules)} candidate rule(s) proposed for "
                      "review.")
    return found


__all__ = ["EXTRACT_VERSION", "Extraction", "availability", "docx_available",
           "extract", "ocr_available", "pdf_available", "rules_of",
           "sections_of", "xlsx_available"]
