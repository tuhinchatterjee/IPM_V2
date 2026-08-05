"""
Pour the report content model into a PDF (reportlab) or a Word file (python-docx).

Both writers walk the same section list and render the same charts, so the two
formats are the same document in two containers. Neither writer computes
anything — if a figure is wrong here it was wrong in content.py.
"""

import io

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from backend.reporting import charts

NAVY = colors.HexColor("#0b2436")
TEAL = colors.HexColor("#16b8a6")
INK = colors.HexColor("#16232f")
MUTED = colors.HexColor("#6c7a8c")
BORDER = colors.HexColor("#e3e8ef")
ROW_ALT = colors.HexColor("#f8fafc")
RED = colors.HexColor("#c0292e")
AMBER = colors.HexColor("#c4690f")

SEVERITY_PDF = {"HIGH": RED, "MEDIUM": AMBER, "LOW": MUTED}
SEVERITY_DOCX = {"HIGH": RGBColor(0xC0, 0x29, 0x2E), "MEDIUM": RGBColor(0xC4, 0x69, 0x0F),
                 "LOW": RGBColor(0x6C, 0x7A, 0x8C)}


# ------------------------------------------------------------------------- PDF

def _pdf_styles():
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("t", parent=base["Title"], fontSize=20, leading=25,
                                textColor=NAVY, alignment=TA_LEFT, spaceAfter=4),
        "subtitle": ParagraphStyle("st", parent=base["Normal"], fontSize=10.5, leading=15,
                                   textColor=MUTED, spaceAfter=2),
        "section": ParagraphStyle("s", parent=base["Heading2"], fontSize=13, leading=17,
                                  textColor=NAVY, spaceBefore=16, spaceAfter=6),
        "body": ParagraphStyle("b", parent=base["Normal"], fontSize=9.5, leading=14,
                               textColor=INK, spaceAfter=8),
        "small": ParagraphStyle("sm", parent=base["Normal"], fontSize=8, leading=11,
                                textColor=MUTED),
        "cell": ParagraphStyle("c", parent=base["Normal"], fontSize=8, leading=11, textColor=INK),
        "cellhead": ParagraphStyle("ch", parent=base["Normal"], fontSize=7.5, leading=10,
                                   textColor=colors.white, fontName="Helvetica-Bold"),
        "finding": ParagraphStyle("f", parent=base["Normal"], fontSize=9, leading=13,
                                  textColor=INK, leftIndent=10, spaceAfter=4),
    }


def _pdf_table(table, styles, col_widths=None):
    head = [Paragraph(str(c), styles["cellhead"]) for c in table["columns"]]
    body = [[Paragraph(str(c), styles["cell"]) for c in row] for row in table["rows"]]
    if not body:
        return None
    t = Table([head] + body, colWidths=col_widths, repeatRows=1, hAlign="LEFT")
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("GRID", (0, 0), (-1, -1), 0.4, BORDER),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, ROW_ALT]),
    ]))
    return t


def _page_furniture(report):
    def draw(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(MUTED)
        canvas.setFont("Helvetica", 7)
        canvas.drawString(18 * mm, 12 * mm, report["classification"])
        canvas.drawRightString(A4[0] - 18 * mm, 12 * mm, f"Page {doc.page}")
        canvas.drawString(18 * mm, A4[1] - 12 * mm,
                          f"{report['short_title']} · {report['quarter_label']}")
        canvas.setStrokeColor(BORDER)
        canvas.setLineWidth(0.4)
        canvas.line(18 * mm, 15 * mm, A4[0] - 18 * mm, 15 * mm)
        canvas.restoreState()
    return draw


def write_pdf(report: dict, context: dict | None = None) -> bytes:
    styles = _pdf_styles()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
        title=f"{report['short_title']} — {report['quarter_label']}",
        author=report["prepared_by"],
    )

    story = [
        Paragraph(report["title"], styles["title"]),
        Paragraph(f"Reporting period {report['quarter_label']} · {report['audience']}",
                  styles["subtitle"]),
        Paragraph(f"Prepared {report['generated_at']} by {report['prepared_by']}", styles["subtitle"]),
        Spacer(1, 8),
        Paragraph(report["purpose"], styles["body"]),
        Spacer(1, 4),
    ]

    contents = [[Paragraph(f"{i}.", styles["cell"]), Paragraph(s["title"], styles["cell"])]
                for i, s in enumerate(report["sections"], 1)]
    toc = Table(contents, colWidths=[12 * mm, 140 * mm], hAlign="LEFT")
    toc.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"),
                             ("TOPPADDING", (0, 0), (-1, -1), 1.5),
                             ("BOTTOMPADDING", (0, 0), (-1, -1), 1.5)]))
    story += [Paragraph("Contents", styles["section"]), toc, PageBreak()]

    for i, section in enumerate(report["sections"], 1):
        story.append(Paragraph(f"{i}. {section['title']}", styles["section"]))
        if section.get("narrative"):
            story.append(Paragraph(section["narrative"], styles["body"]))

        png = charts.render(section.get("chart"), context)
        if png:
            img = Image(io.BytesIO(png))
            avail = A4[0] - 36 * mm
            scale = min(avail / img.imageWidth, 1.0)
            img.drawWidth = img.imageWidth * scale
            img.drawHeight = img.imageHeight * scale
            story += [Spacer(1, 2), img, Spacer(1, 8)]

        for key in ("table", "extra_table"):
            t = section.get(key)
            if t and t.get("rows"):
                built = _pdf_table(t, styles)
                if built is not None:
                    story += [built, Spacer(1, 8)]

        findings = section.get("findings") or []
        if findings:
            story.append(Paragraph("Findings", ParagraphStyle(
                "fh", parent=styles["body"], fontName="Helvetica-Bold", spaceAfter=3)))
            for f in findings:
                story.append(Paragraph(
                    f'<font color="{SEVERITY_PDF[f["severity"]].hexval()}"><b>[{f["severity"]}]</b></font> '
                    f'{f["text"]}', styles["finding"]))
            story.append(Spacer(1, 6))

    story += [
        PageBreak(),
        Paragraph("Basis of preparation", styles["section"]),
        Paragraph(
            "Every figure in this pack is computed from the active portfolio dataset at the stated "
            "reporting date and reconciles to the corresponding screen in the IPM tool. Capital-linked "
            "figures use a documented risk-weight proxy; macro scenario paths and the climate stressed-PD "
            "model are disclosed as model output rather than measured outcomes. Recommended actions and "
            "remediation items are derived from the findings raised in this pack.",
            styles["body"]),
        Paragraph(report["classification"], styles["small"]),
    ]

    furniture = _page_furniture(report)
    doc.build(story, onFirstPage=furniture, onLaterPages=furniture)
    return buf.getvalue()


# ------------------------------------------------------------------------ DOCX

def _docx_shade(cell, hex_fill):
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shade)


def _docx_table(document, table):
    if not table.get("rows"):
        return
    t = document.add_table(rows=1, cols=len(table["columns"]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for cell, name in zip(t.rows[0].cells, table["columns"], strict=True):
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(name))
        run.bold = True
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _docx_shade(cell, "0B2436")
    for row in table["rows"]:
        cells = t.add_row().cells
        for cell, value in zip(cells, row, strict=True):
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(value))
            run.font.size = Pt(8)
    return t


def write_docx(report: dict, context: dict | None = None) -> bytes:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.text = f"{report['short_title']} · {report['quarter_label']}"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(0x6C, 0x7A, 0x8C)
    footer = section.footer.paragraphs[0]
    footer.text = report["classification"]
    footer.runs[0].font.size = Pt(7)
    footer.runs[0].font.color.rgb = RGBColor(0x6C, 0x7A, 0x8C)

    title = document.add_heading(report["title"], level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x24, 0x36)

    meta = document.add_paragraph()
    meta.add_run(f"Reporting period {report['quarter_label']} · {report['audience']}\n").font.size = Pt(9.5)
    meta.add_run(f"Prepared {report['generated_at']} by {report['prepared_by']}").font.size = Pt(9.5)
    for run in meta.runs:
        run.font.color.rgb = RGBColor(0x6C, 0x7A, 0x8C)

    document.add_paragraph(report["purpose"])

    document.add_heading("Contents", level=1)
    for i, s in enumerate(report["sections"], 1):
        p = document.add_paragraph(f"{i}.  {s['title']}")
        p.paragraph_format.space_after = Pt(2)

    document.add_section(WD_SECTION.NEW_PAGE)

    for i, sec in enumerate(report["sections"], 1):
        document.add_heading(f"{i}. {sec['title']}", level=1)
        if sec.get("narrative"):
            document.add_paragraph(sec["narrative"])

        png = charts.render(sec.get("chart"), context)
        if png:
            document.add_picture(io.BytesIO(png), width=Inches(6.2))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT

        for key in ("table", "extra_table"):
            t = sec.get(key)
            if t and t.get("rows"):
                _docx_table(document, t)
                document.add_paragraph()

        findings = sec.get("findings") or []
        if findings:
            document.add_heading("Findings", level=2)
            for f in findings:
                p = document.add_paragraph(style="List Bullet")
                tag = p.add_run(f"[{f['severity']}] ")
                tag.bold = True
                tag.font.color.rgb = SEVERITY_DOCX[f["severity"]]
                p.add_run(f["text"])

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Basis of preparation", level=1)
    document.add_paragraph(
        "Every figure in this pack is computed from the active portfolio dataset at the stated "
        "reporting date and reconciles to the corresponding screen in the IPM tool. Capital-linked "
        "figures use a documented risk-weight proxy; macro scenario paths and the climate stressed-PD "
        "model are disclosed as model output rather than measured outcomes. Recommended actions and "
        "remediation items are derived from the findings raised in this pack.")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


FORMATS = {
    "pdf": {"label": "PDF", "extension": "pdf", "writer": write_pdf,
            "mime": "application/pdf"},
    "docx": {"label": "Word", "extension": "docx", "writer": write_docx,
             "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
}


def write(report: dict, fmt: str = "pdf", context: dict | None = None) -> tuple[bytes, str, str]:
    """(bytes, filename, mime) for one report in one format."""
    spec = FORMATS.get(fmt, FORMATS["pdf"])
    data = spec["writer"](report, context)
    stamp = report["quarter"].replace(" ", "")
    name = f"{report['type'].upper()}_{stamp}_Report.{spec['extension']}"
    return data, name, spec["mime"]
