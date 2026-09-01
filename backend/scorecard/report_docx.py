"""
The validation report as a Word document. §51, §53.

This file makes no decisions about content. Every figure, caption, table and
refusal was settled in `report.py`; this pours that into a .docx and worries
only about how it looks on a page. If a number here is wrong it was wrong
before this file saw it, which is the property that lets §51's "do not ask
the model to recalculate numbers" hold trivially: nothing here calculates.

§53 asks for a top-tier model-risk look. Concretely that means a title
page, a document-control table, a table of contents, running headers and
footers with page numbers, numbered sections, consistent table styling,
captions on every table, and the disclaimer where a reader cannot miss it.

The table of contents is a field, not a list
----------------------------------------------
Word builds it on open (or on F9). Writing the page numbers ourselves would
mean guessing where the pagination falls, and a contents page whose numbers
are wrong is worse than none: it tells a reader the document was assembled
by something that did not read it.
"""

from __future__ import annotations

import io

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from backend.scorecard import report as report_mod

INK = RGBColor(0x16, 0x23, 0x2F)
MUTED = RGBColor(0x6C, 0x7A, 0x8C)
NAVY = "0B2436"
BAND = "F4F6F8"

#: Severity colours for the findings table. Grey for an observation: an
#: observation rendered in red is a finding, whatever the column says.
SEVERITY_INK = {
    "HIGH": RGBColor(0xC0, 0x29, 0x2E),
    "MEDIUM": RGBColor(0xC4, 0x69, 0x0F),
    "LOW": RGBColor(0x6C, 0x7A, 0x8C),
    "OBSERVATION": RGBColor(0x6C, 0x7A, 0x8C),
}


def _shade(cell, fill: str) -> None:
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(element)


def _field(paragraph, instruction: str) -> None:
    """Insert a Word field, for the page number and the contents."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, text, separate, end):
        run._r.append(element)


def _table(document, table: report_mod.Table) -> None:
    """One captioned table. Empty tables are skipped, not drawn hollow."""
    if not table.rows:
        return
    caption = document.add_paragraph()
    run = caption.add_run(table.caption)
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = INK
    caption.paragraph_format.space_after = Pt(2)

    grid = document.add_table(rows=1, cols=len(table.columns))
    grid.style = "Table Grid"
    grid.alignment = WD_TABLE_ALIGNMENT.LEFT
    for cell, name in zip(grid.rows[0].cells, table.columns, strict=True):
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(name))
        run.bold = True
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(cell, NAVY)

    for index, row in enumerate(table.rows):
        cells = grid.add_row().cells
        for position, value in enumerate(row):
            cell = cells[position] if position < len(cells) else None
            if cell is None:
                continue
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(value))
            run.font.size = Pt(8)
            if position == 0 and str(value).upper() in SEVERITY_INK:
                run.bold = True
                run.font.color.rgb = SEVERITY_INK[str(value).upper()]
            if index % 2 == 1:
                _shade(cell, BAND)

    if table.note:
        note = document.add_paragraph()
        run = note.add_run(table.note)
        run.italic = True
        run.font.size = Pt(7.5)
        run.font.color.rgb = MUTED
    document.add_paragraph()


def _title_page(document, report: report_mod.Report) -> None:
    for _ in range(3):
        document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Model Validation Report")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = INK

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(report.title)
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED

    stamp = document.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stamp.add_run(
        f"{report.model_name} · version {report.model_version}\n"
        f"Validation period {report.period}\n"
        f"Report {report.report_id} · generated {report.generated_at}")
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    opinion = document.add_paragraph()
    opinion.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = opinion.add_run(f"\nValidation opinion: {report.opinion}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = INK

    for _ in range(2):
        document.add_paragraph()

    # §2, on the title page rather than in a footnote. Somebody who reads
    # only the cover has to know what the numbers describe.
    notice = document.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = notice.add_run(report_mod.SYNTHETIC_NOTICE)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    disclaimer = document.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(report.disclaimer)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def write(report: report_mod.Report) -> bytes:
    """The report, as .docx bytes."""
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.color.rgb = INK

    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.text = (f"{report.scorecard_type.title()} Scorecard Validation · "
                   f"{report.period} · {report.model_version}")
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"{report.report_id} · Internal — model validation"
                         f" · {report.origin} · Page ")
    run.font.size = Pt(7)
    run.font.color.rgb = MUTED
    _field(footer, "PAGE")

    _title_page(document, report)

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Contents", level=1)
    contents = document.add_paragraph()
    _field(contents, r'TOC \o "1-2" \h \z \u')
    note = document.add_paragraph()
    run = note.add_run(
        "If the contents above are blank, press F9 in Word to build them.")
    run.italic = True
    run.font.size = Pt(7.5)
    run.font.color.rgb = MUTED

    document.add_section(WD_SECTION.NEW_PAGE)
    for entry in report.sections:
        heading = document.add_heading(
            f"{entry.number} {entry.title}", level=min(entry.level, 3))
        for run in heading.runs:
            run.font.color.rgb = INK
        if entry.narrative:
            document.add_paragraph(entry.narrative)
        if entry.unavailable:
            # Never a zero. §7 and §50 both come down to this paragraph.
            paragraph = document.add_paragraph()
            run = paragraph.add_run("Not reported. ")
            run.bold = True
            run.font.size = Pt(9)
            run = paragraph.add_run(entry.unavailable)
            run.font.size = Pt(9)
            run.font.color.rgb = MUTED
        for table in entry.tables:
            _table(document, table)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
