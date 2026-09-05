"""The lifecycle over HTTP, through the real app.

The point of this file is not to re-test the service — the service has its own
suite. It is to check the three things a router can get wrong on its own:

  * a refusal arriving as the wrong status code, so a screen shows "something
    went wrong" where it should show "you need approver access"
  * the source being readable from the request, which would let a caller
    launder an agent write into a human one
  * a download arriving as JSON rather than as bytes with a filename
"""

from __future__ import annotations

import uuid

import pytest

from tests.conftest import database_available

pytestmark = pytest.mark.skipif(
    not database_available(), reason="PostgreSQL not reachable")


@pytest.fixture(scope="module")
def client():
    from fastapi.testclient import TestClient

    from backend.api.main import create_app

    with TestClient(create_app()) as handle:
        yield handle


def _as(user_id: int | None, role: str = "ANALYST") -> dict[str, str]:
    """Act as one person. The documented header path for a deployment with
    signing in switched off — see tests/conftest.py."""
    headers = {"X-IPM-Role": role}
    if user_id is not None:
        headers["X-IPM-User-Id"] = str(user_id)
    return headers


@pytest.fixture
def committee_over_http(client, people, session):
    """A committee created through the API, torn down through the ORM."""
    session.commit()  # the API opens its own session; these people must exist
    made = client.post(
        "/api/v1/playbook/committees",
        headers=_as(int(people["steward"].id), "DATA_STEWARD"),
        json={"name": f"HTTP Committee {uuid.uuid4().hex[:6]}",
              "business_area": "Retail Credit Risk", "cadence": "MONTHLY",
              "meeting_weekday": 2})
    assert made.status_code == 201, made.text
    body = made.json()

    for key, level in (("owner", "OWNER"), ("author", "CONTRIBUTOR"),
                       ("reviewer", "REVIEWER"), ("approver", "APPROVER")):
        added = client.post(
            f"/api/v1/playbook/committees/{body['id']}/members",
            headers=_as(int(people["steward"].id), "DATA_STEWARD"),
            json={"user_id": int(people[key].id), "access_role": level})
        assert added.status_code == 201, added.text

    yield body

    from backend.models.playbook import PlaybookCommittee

    row = session.get(PlaybookCommittee, int(body["id"]))
    if row is not None:
        session.delete(row)
    session.commit()


@pytest.fixture
def pack_over_http(client, committee_over_http, people):
    made = client.post(
        "/api/v1/playbook/packs", headers=_as(int(people["owner"].id)),
        json={"committee_id": committee_over_http["id"], "period": "2025-01",
              "comparison_period": "2024-12"})
    assert made.status_code == 201, made.text
    return made.json()


# ============================================================ the lifecycle


def test_a_pack_can_be_created_read_and_generated_over_http(
        client, pack_over_http, people):
    who = _as(int(people["owner"].id))
    pack_id = pack_over_http["id"]

    section = client.post(
        f"/api/v1/playbook/packs/{pack_id}/sections", headers=who,
        json={"title": "Portfolio performance"})
    assert section.status_code == 201, section.text

    block = client.post(
        f"/api/v1/playbook/sections/{section.json()['id']}/blocks",
        headers=who,
        json={"block_type": "KPI", "title": "Retail default rate",
              "config": {"metric_id": "retail.default_rate"}})
    assert block.status_code == 201, block.text

    run = client.post(f"/api/v1/playbook/packs/{pack_id}/generate",
                      headers=who)
    assert run.status_code == 200, run.text
    assert run.json()["calculated"] == 1

    whole = client.get(f"/api/v1/playbook/packs/{pack_id}", headers=who)
    assert whole.status_code == 200
    figure = whole.json()["sections"][0]["blocks"][0]["figure"]
    assert figure is not None
    assert figure["formula_hash"], "the working travels with the number"


def test_readiness_over_http_says_what_is_blocking(client, pack_over_http,
                                                   people):
    who = _as(int(people["owner"].id))
    client.post(f"/api/v1/playbook/packs/{pack_over_http['id']}/sections",
                headers=who, json={"title": "A required section"})

    state = client.get(
        f"/api/v1/playbook/packs/{pack_over_http['id']}/readiness",
        headers=who)
    assert state.status_code == 200
    body = state.json()
    assert body["state"] in ("RED", "AMBER", "GREEN")
    assert "computed_at" in body, (
        "a percentage with no timestamp is a number nobody can defend")
    assert isinstance(body["reasons"], list)


# ========================================================== the status codes


def test_a_pack_somebody_may_not_see_is_a_404_not_a_403(client, pack_over_http,
                                                        people):
    answer = client.get(f"/api/v1/playbook/packs/{pack_over_http['id']}",
                        headers=_as(int(people["outsider"].id)))
    assert answer.status_code == 404
    assert answer.json()["detail"]["error"] == "not_found"


def test_a_pack_that_does_not_exist_answers_identically(client, people):
    answer = client.get("/api/v1/playbook/packs/2000000001",
                        headers=_as(int(people["outsider"].id)))
    assert answer.status_code == 404
    assert answer.json()["detail"]["error"] == "not_found"


def test_too_little_access_is_a_403_that_says_what_is_needed(
        client, pack_over_http, people):
    """A screen has to be able to say "you need approver access"."""
    answer = client.post(
        f"/api/v1/playbook/packs/{pack_over_http['id']}/status",
        headers=_as(int(people["author"].id)),
        json={"status": "CONTRIBUTOR_REVIEW"})
    assert answer.status_code == 403, answer.text
    detail = answer.json()["detail"]
    assert detail["error"] == "forbidden"
    assert "access" in detail["message"]


def test_a_locked_pack_is_a_409_and_not_a_403(client, pack_over_http, people,
                                              session):
    """Denied means ask for access; locked means raise an amendment."""
    from backend.models.playbook import PlaybookPack

    row = session.get(PlaybookPack, int(pack_over_http["id"]))
    row.status = "APPROVED"
    session.commit()

    answer = client.patch(
        f"/api/v1/playbook/packs/{pack_over_http['id']}",
        headers=_as(int(people["owner"].id)), json={"name": "Changed"})
    assert answer.status_code == 409, answer.text
    detail = answer.json()["detail"]
    assert detail["error"] == "pack_locked"
    assert "amendment" in detail["message"]


def test_a_stale_write_is_a_409_that_names_who_moved_it(client, pack_over_http,
                                                        people):
    who = _as(int(people["owner"].id))
    version = pack_over_http["version"]
    first = client.patch(f"/api/v1/playbook/packs/{pack_over_http['id']}",
                         headers=who,
                         json={"expected_version": version,
                               "name": "Their edit"})
    assert first.status_code == 200, first.text

    second = client.patch(f"/api/v1/playbook/packs/{pack_over_http['id']}",
                          headers=who,
                          json={"expected_version": version,
                                "name": "My edit"})
    assert second.status_code == 409
    detail = second.json()["detail"]
    assert detail["error"] == "stale_write"
    assert "Owner Tester" in detail["message"]


def test_an_unrepresentable_request_is_a_422_naming_the_vocabulary(
        client, pack_over_http, people):
    answer = client.post(
        f"/api/v1/playbook/packs/{pack_over_http['id']}/status",
        headers=_as(int(people["owner"].id)), json={"status": "SPLENDID"})
    assert answer.status_code == 422
    assert "One of:" in answer.json()["detail"]["message"]


# ================================================================ the source


def test_the_source_cannot_be_named_by_the_caller(client, pack_over_http,
                                                  people, session):
    """A caller who could name their own source could name UI.

    The router passes SOURCE_UI on every call and never reads one from the
    request, so a body carrying `source` is simply not part of the contract —
    and the write that lands is recorded as UI, not as whatever was sent.
    """
    who = _as(int(people["owner"].id))
    answer = client.patch(
        f"/api/v1/playbook/packs/{pack_over_http['id']}", headers=who,
        json={"name": "Renamed", "source": "SYSTEM"})
    assert answer.status_code == 200, answer.text

    from sqlalchemy import select

    from backend.models.playbook import PlaybookEvent

    session.expire_all()
    event = session.execute(
        select(PlaybookEvent).where(
            PlaybookEvent.pack_id == pack_over_http["id"],
            PlaybookEvent.action == "updated")
        .order_by(PlaybookEvent.id.desc())).scalars().first()
    assert event is not None
    assert event.source == "UI", (
        "an extra field in the body must not decide how the change is "
        "recorded")


def test_a_caller_cannot_claim_a_block_was_imported(client, pack_over_http,
                                                    people):
    """`import_class` is the one field that waives the metric rule.

    An UNMAPPED_TABLE is a table CreditProbe did not calculate, and it is the
    only calculated block allowed to name no metric. If a caller could set
    that field, they could put an unlabelled table of their own numbers on a
    committee pack and it would look like everything around it. The router
    does not accept the field at all, so the rule still bites.
    """
    who = _as(int(people["owner"].id))
    section = client.post(
        f"/api/v1/playbook/packs/{pack_over_http['id']}/sections",
        headers=who, json={"title": "Portfolio"})

    answer = client.post(
        f"/api/v1/playbook/sections/{section.json()['id']}/blocks",
        headers=who,
        json={"block_type": "TABLE", "title": "Numbers I typed",
              "import_class": "UNMAPPED_TABLE",
              "config": {"rows": [["Mass", "0.1"]]}})
    assert answer.status_code == 422, answer.text
    assert "name the metric" in answer.json()["detail"]["message"]


def test_an_oversized_upload_is_refused_without_being_taken_whole(
        client, pack_over_http, people):
    """The limit has to cost what the limit costs.

    A body far over the limit must be refused after the server has read a
    little past it, not after it has accepted all of it — otherwise the limit
    protects the parser and nothing else.
    """
    from backend.playbook import import_ as ingest

    over = b"PK\x03\x04" + b"x" * (ingest.MAX_BYTES + 4096)
    answer = client.post(
        f"/api/v1/playbook/packs/{pack_over_http['id']}/import",
        headers=_as(int(people["owner"].id)),
        files={"file": ("huge.docx", over,
                        "application/vnd.openxmlformats-officedocument."
                        "wordprocessingml.document")})
    assert answer.status_code == 413, answer.text
    detail = answer.json()["detail"]
    assert detail["error"] == "too_large"
    assert "Data Builder" in detail["message"], (
        "the refusal says where a file that size does belong")


def test_an_import_over_http_labels_what_it_produced(client, pack_over_http,
                                                     people):
    import io

    from docx import Document

    document = Document()
    document.add_heading("Their portfolio section", level=1)
    document.add_paragraph(
        "Balances grew by four per cent over the quarter under review.")
    buffer = io.BytesIO()
    document.save(buffer)

    who = _as(int(people["owner"].id))
    answer = client.post(
        f"/api/v1/playbook/packs/{pack_over_http['id']}/import", headers=who,
        files={"file": ("theirs.docx", buffer.getvalue(),
                        "application/vnd.openxmlformats-officedocument."
                        "wordprocessingml.document")})
    assert answer.status_code == 201, answer.text
    assert answer.json()["source_id"]

    whole = client.get(f"/api/v1/playbook/packs/{pack_over_http['id']}",
                       headers=who).json()
    imported = [b for s in whole["sections"] for b in s["blocks"]
                if b["import_class"]]
    assert imported
    assert all(b["source"] == "IMPORT" for b in imported)


# ================================================================ downloads


def test_a_download_is_bytes_with_a_filename_and_a_checksum(
        client, pack_over_http, people):
    who = _as(int(people["owner"].id))
    pack_id = pack_over_http["id"]

    section = client.post(f"/api/v1/playbook/packs/{pack_id}/sections",
                          headers=who, json={"title": "Portfolio"})
    client.post(f"/api/v1/playbook/sections/{section.json()['id']}/blocks",
                headers=who,
                json={"block_type": "KPI", "title": "Default rate",
                      "config": {"metric_id": "retail.default_rate"}})
    client.post(f"/api/v1/playbook/packs/{pack_id}/generate", headers=who)

    answer = client.get(f"/api/v1/playbook/packs/{pack_id}/export?format=pdf",
                        headers=who)
    assert answer.status_code == 200, answer.text
    assert answer.headers["content-type"] == "application/pdf"
    assert "attachment; filename=" in answer.headers["content-disposition"]
    assert len(answer.headers["x-creditprobe-checksum"]) == 64
    assert answer.content.startswith(b"%PDF-")


def test_every_offered_format_actually_downloads(client, pack_over_http,
                                                 people):
    who = _as(int(people["owner"].id))
    pack_id = pack_over_http["id"]
    section = client.post(f"/api/v1/playbook/packs/{pack_id}/sections",
                          headers=who, json={"title": "Portfolio"})
    client.post(f"/api/v1/playbook/sections/{section.json()['id']}/blocks",
                headers=who,
                json={"block_type": "KPI", "title": "Default rate",
                      "config": {"metric_id": "retail.default_rate"}})
    client.post(f"/api/v1/playbook/packs/{pack_id}/generate", headers=who)

    offered = client.get("/api/v1/playbook/formats", headers=who).json()
    assert offered["formats"], "the button offers something"

    for entry in offered["formats"]:
        answer = client.get(
            f"/api/v1/playbook/packs/{pack_id}/export?format={entry['format']}",
            headers=who)
        assert answer.status_code == 200, (entry["format"], answer.text)
        assert len(answer.content) > 1000, entry["format"]


def test_an_unknown_download_format_is_refused(client, pack_over_http, people):
    answer = client.get(
        f"/api/v1/playbook/packs/{pack_over_http['id']}/export?format=exe",
        headers=_as(int(people["owner"].id)))
    assert answer.status_code == 422
    assert "not a format" in answer.json()["detail"]["message"]


def test_somebody_who_cannot_read_a_pack_cannot_download_it(
        client, pack_over_http, people):
    answer = client.get(
        f"/api/v1/playbook/packs/{pack_over_http['id']}/export?format=pdf",
        headers=_as(int(people["outsider"].id)))
    assert answer.status_code == 404


# ================================================================= findings


def test_findings_reach_a_screen_and_can_be_answered_over_http(
        client, pack_over_http, people, session):
    """Generation raises findings; without these routes nobody could see one."""
    who = _as(int(people["owner"].id))
    pack_id = pack_over_http["id"]

    from backend.models.playbook import PlaybookFinding

    session.add(PlaybookFinding(
        pack_id=int(pack_id), finding_type="THRESHOLD_BREACH", severity="HIGH",
        title="Default rate above the agreed band",
        factual_basis="Observed 6.88% against a ceiling of 6.50%.",
        rule_key="default_rate_band", fingerprint="http-findings-test",
        status="OPEN"))
    session.commit()

    listed = client.get(f"/api/v1/playbook/findings?pack_id={pack_id}",
                        headers=who)
    assert listed.status_code == 200, listed.text
    raised = listed.json()["findings"]
    assert raised, "a finding on a pack is visible to the pack's owner"
    finding_id = raised[0]["id"]
    assert raised[0]["rule_key"], "the rule that fired travels to the screen"

    answered = client.post(
        f"/api/v1/playbook/findings/{finding_id}/respond", headers=who,
        json={"status": "EXPLAINED",
              "response": "Two 2024 vintages, already being reworked."})
    assert answered.status_code == 200, answered.text
    assert answered.json()["answered"] is True


def test_a_dismissal_over_http_without_a_reason_is_a_422(client,
                                                         pack_over_http,
                                                         people, session):
    from backend.models.playbook import PlaybookFinding

    session.add(PlaybookFinding(
        pack_id=int(pack_over_http["id"]), finding_type="DATA_QUALITY",
        severity="MEDIUM", title="A gap worth explaining",
        factual_basis="Rows missing for one segment.",
        fingerprint="http-dismissal-test", status="OPEN"))
    session.commit()

    who = _as(int(people["approver"].id))
    listed = client.get(
        f"/api/v1/playbook/findings?pack_id={pack_over_http['id']}",
        headers=who).json()["findings"]
    finding_id = next(f["id"] for f in listed
                      if f["title"] == "A gap worth explaining")

    answer = client.post(
        f"/api/v1/playbook/findings/{finding_id}/respond", headers=who,
        json={"status": "DISMISSED"})
    assert answer.status_code == 422, answer.text
    assert "needs a reason" in answer.json()["detail"]["message"]

    # And with one, the same call goes through — so the 422 above is the
    # missing reason and not a broken route.
    with_reason = client.post(
        f"/api/v1/playbook/findings/{finding_id}/respond", headers=who,
        json={"status": "DISMISSED",
              "reason": "The segment was retired in December."})
    assert with_reason.status_code == 200, with_reason.text
    assert with_reason.json()["dismissed_by"] == int(people["approver"].id)


def test_a_finding_id_from_another_committee_is_a_404(client, pack_over_http,
                                                      people, session):
    from backend.models.playbook import PlaybookFinding

    session.add(PlaybookFinding(
        pack_id=int(pack_over_http["id"]), finding_type="DATA_QUALITY",
        severity="LOW", title="Not yours to read",
        factual_basis="Nothing much.", fingerprint="http-idor-test",
        status="OPEN"))
    session.commit()
    listed = client.get(
        f"/api/v1/playbook/findings?pack_id={pack_over_http['id']}",
        headers=_as(int(people["owner"].id))).json()["findings"]

    answer = client.get(f"/api/v1/playbook/findings/{listed[0]['id']}",
                        headers=_as(int(people["outsider"].id)))
    assert answer.status_code == 404


# ========================================================= removing a page


def test_a_section_added_over_http_can_be_removed_over_http(client,
                                                            pack_over_http,
                                                            people):
    who = _as(int(people["owner"].id))
    made = client.post(
        f"/api/v1/playbook/packs/{pack_over_http['id']}/sections",
        headers=who, json={"title": "Added then removed"})
    assert made.status_code == 201, made.text

    gone = client.delete(f"/api/v1/playbook/sections/{made.json()['id']}",
                         headers=who)
    assert gone.status_code == 204

    whole = client.get(f"/api/v1/playbook/packs/{pack_over_http['id']}",
                       headers=who).json()
    assert "Added then removed" not in [s["title"] for s in whole["sections"]]


# ==================================================================== chase


def test_the_chase_screen_delivers_nothing(client, committee_over_http,
                                           people, session):
    """The screen a pack owner opens must not notify everyone it names."""
    from sqlalchemy import func, select

    from backend.models.playbook import PlaybookReminder

    before = int(session.execute(
        select(func.count()).select_from(PlaybookReminder)).scalar_one())
    answer = client.get(
        f"/api/v1/playbook/chase?committee_id={committee_over_http['id']}",
        headers=_as(int(people["owner"].id)))
    assert answer.status_code == 200, answer.text
    assert "outstanding" in answer.json()

    session.expire_all()
    after = int(session.execute(
        select(func.count()).select_from(PlaybookReminder)).scalar_one())
    assert after == before, "a dry run writes nothing and sends nothing"
