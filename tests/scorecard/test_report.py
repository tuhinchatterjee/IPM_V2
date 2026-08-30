"""
The CBUAE-aligned validation report and its evidence workbook.
§51-§56, §83, §89.

Two things are being tested and they are different.

The first is coverage: §89 lists seventeen topics a validation report has to
address, and `report.COVERAGE` maps each to a section. The test asserts the
section exists *and has content*, because a report that contains the word
"calibration" under an empty heading has addressed nothing.

The second is refusal. A report over a month whose performance window has
not closed must say so, in the outcome sections, in place of numbers — not
alongside them and not as zeros. That is §7, and it is the property most
likely to be quietly lost the next time somebody makes the report "more
complete".
"""

from __future__ import annotations

import io

import pytest
from docx import Document
from openpyxl import load_workbook

from backend.scorecard import dashboard as dash
from backend.scorecard import report as rpt
from backend.scorecard import report_docx, report_xlsx
from backend.scorecard import synthetic as synth

APP = "APPLICATION"
BEH = "BEHAVIORAL"


@pytest.fixture(scope="module")
def matured_month() -> str:
    months = dash.available_months(APP)
    return [m for m in months if synth.matured(m)][-1]


@pytest.fixture(scope="module")
def immature_month() -> str:
    months = dash.available_months(APP)
    open_windows = [m for m in months if not synth.matured(m)]
    assert open_windows, (
        "the demonstration universe has no month whose performance window is "
        "still open, so §7's refusal cannot be reached from any screen")
    return open_windows[0]


@pytest.fixture(scope="module")
def report(matured_month):
    return rpt.build(APP, month=matured_month, generated_by="ADMIN#1")


@pytest.fixture(scope="module")
def open_report(immature_month):
    return rpt.build(APP, month=immature_month, generated_by="ADMIN#1")


# ------------------------------------------------------------- §89 coverage


def test_the_report_addresses_every_topic_the_standard_lists(report):
    """§89. Seventeen topics, each in a section that has content."""
    result = rpt.coverage(report)
    assert result["missing"] == []
    assert result["complete"] is True
    assert result["topics"] == len(rpt.COVERAGE)


def test_every_coverage_target_names_a_section_that_exists(report):
    """A coverage map pointing at a section number nobody wrote would report
    a hole as filled."""
    numbers = {s.number for s in report.sections}
    for topic, number in rpt.COVERAGE.items():
        assert number in numbers, f"{topic} points at missing section {number}"


def test_coverage_notices_an_empty_section(report):
    """The check has to be able to fail, or it is decoration."""
    hollowed = rpt.Report(
        report_id=report.report_id, model_id=report.model_id,
        model_version=report.model_version, model_name=report.model_name,
        scorecard_type=report.scorecard_type, model_kind=report.model_kind,
        period=report.period, title=report.title,
        structure_version=report.structure_version,
        generated_at=report.generated_at, generated_by=report.generated_by,
        opinion=report.opinion,
        sections=[rpt.Section("8.3", "Calibration and accuracy")])
    result = rpt.coverage(hollowed)
    assert result["complete"] is False
    assert "calibration" in result["missing"]


def test_the_thirteen_top_level_sections_are_present(report):
    top = [s.number for s in report.sections if s.level == 1]
    assert top == [str(n) for n in range(1, 14)]


# ------------------------------------------------------------------ §7


def test_an_open_window_refuses_the_outcome_sections(open_report,
                                                     immature_month):
    """§7. Never actual against predicted on an immature cohort."""
    for number in ("8.2", "8.3"):
        section = open_report.section(number)
        assert section.unavailable, f"{number} reported an outcome metric"
        assert section.tables == [], (
            f"{number} is unavailable and still drew a table, which is how a "
            "row of dashes gets read as a measurement")
        assert immature_month in section.unavailable


def test_an_open_window_still_reports_stability(open_report):
    """Stability needs no outcome, so withholding it would be a different
    kind of dishonesty: refusing to show what is actually known."""
    stability = open_report.section("8.4")
    assert not stability.unavailable
    assert stability.tables


def test_the_refusal_says_when_the_window_closes(open_report,
                                                 immature_month):
    """"Not available" without a date reads as broken rather than honest."""
    closes = synth.window_closes(immature_month)
    assert closes in open_report.section("8.3").unavailable


def test_the_default_month_is_the_latest_matured_one(matured_month):
    """§18. Not the latest month — the latest month with an answer."""
    assert rpt.build(APP, generated_by="t").period == matured_month


# ------------------------------------------------- §0 and §2, the two claims


def test_the_report_never_claims_a_certification(report):
    assert "does not provide regulatory certification" in report.disclaimer
    conclusion = report.section("12")
    assert "certification" in conclusion.narrative.lower()


def test_the_report_says_its_data_is_synthetic(report):
    assert report.origin == synth.ORIGIN
    assert "no real customer" in rpt.SYNTHETIC_NOTICE
    assert rpt.SYNTHETIC_NOTICE in report.section("1").narrative


def test_no_limit_is_presented_without_its_source(report):
    """§26/§80. The limits table's Source column, and the note under it."""
    monitoring = report.section("9")
    limits = monitoring.tables[0]
    assert limits.columns == ["Metric", "Observed", "Limit", "Status",
                              "Source"]
    assert "not regulatory requirements" in limits.note
    for row in limits.rows:
        assert row[4], f"{row[0]} has a status with no limit source"


def test_the_regulatory_mapping_maps_structure_not_thresholds():
    """§89's last line: do not claim specific regulatory metric limits
    without sourced support. The mapping names sections, not rules."""
    for topic, section in report_xlsx.REGULATORY_MAPPING:
        assert topic and section
        assert section[0].isdigit(), (
            f"{topic} maps to {section!r}, which is not a section of this "
            "report")


# --------------------------------------------------------------- §55 index


def test_every_evidence_item_can_be_found_again(report):
    """§55. Period, model version, method, validation state and the sheet."""
    assert report.evidence
    sheets = set(report_xlsx.SHEETS)
    for item in report.evidence:
        assert item.period == report.period
        assert item.model_version
        assert item.method
        assert item.validation_state
        assert item.workbook_sheet in sheets, (
            f"{item.label} points at sheet {item.workbook_sheet!r}, which the "
            "workbook does not have")


def test_the_evidence_index_appears_in_the_report_itself(report):
    index = report.section("13.5")
    assert index is not None
    assert index.tables
    assert len(index.tables[0].rows) == len(report.evidence)


# ---------------------------------------------------------------- §56 hash


def test_the_content_hash_ignores_who_generated_it_and_when(matured_month):
    """§56's hash answers "did anything change?", not "was this run twice?".

    Two people regenerating the same month's report get the same hash. The
    cover differs — it names them and the date — and that is exactly why the
    cover is outside the hash.
    """
    first = rpt.build(APP, month=matured_month, generated_by="ADMIN#1",
                      generated_at="2026-01-01T00:00:00+00:00")
    second = rpt.build(APP, month=matured_month, generated_by="ANALYST#9",
                       generated_at="2026-06-30T12:00:00+00:00")
    assert first.content_hash == second.content_hash


def test_a_different_month_hashes_differently(matured_month):
    months = [m for m in dash.available_months(APP) if synth.matured(m)]
    first = rpt.build(APP, month=months[-1], generated_by="t")
    other = rpt.build(APP, month=months[-2], generated_by="t")
    assert first.content_hash != other.content_hash


def test_the_filename_follows_the_named_pattern(report):
    """§51's exact shape."""
    name = rpt.filename_for(report, "docx")
    assert name.startswith("CreditProbe_APPLICATION_INCUMBENT_")
    assert name.endswith("_Validation_Report.docx")
    assert report.period in name


# ------------------------------------------------------------ §53 the DOCX


def test_the_word_report_has_a_heading_for_every_section(report):
    body = report_docx.write(report)
    document = Document(io.BytesIO(body))
    headings = [p.text for p in document.paragraphs
                if p.style.name.startswith("Heading")]
    for section in report.sections:
        assert f"{section.number} {section.title}" in headings


def test_the_word_report_carries_its_furniture(report):
    """§53: a header, a footer with a page-number field, and contents."""
    body = report_docx.write(report)
    document = Document(io.BytesIO(body))
    section = document.sections[0]
    assert report.period in section.header.paragraphs[0].text
    assert report.report_id in section.footer.paragraphs[0].text
    # The two fields live in different parts of the package: the contents
    # in the body, the page number in the footer.
    assert "TOC" in document.element.xml, "no table of contents field"
    assert "PAGE" in section.footer._element.xml, "no page-number field"


def test_the_word_report_puts_the_disclaimer_on_the_cover(report):
    body = report_docx.write(report)
    document = Document(io.BytesIO(body))
    front = " ".join(p.text for p in document.paragraphs[:24])
    assert "does not provide regulatory certification" in front
    assert "no real customer" in front


def test_an_unavailable_section_prints_a_reason_not_a_table(open_report):
    body = report_docx.write(open_report)
    document = Document(io.BytesIO(body))
    prose = " ".join(p.text for p in document.paragraphs)
    assert "Not reported." in prose
    assert open_report.section("8.3").unavailable in prose


# ------------------------------------------------------------- §83 the XLSX


def test_the_workbook_has_every_sheet_the_brief_names(report):
    body = report_xlsx.write(report)
    book = load_workbook(io.BytesIO(body))
    assert book.sheetnames == list(report_xlsx.SHEETS)


def test_the_evidence_index_is_the_first_sheet(report):
    body = report_xlsx.write(report)
    book = load_workbook(io.BytesIO(body))
    assert book.sheetnames[0] == "EVIDENCE INDEX"
    sheet = book["EVIDENCE INDEX"]
    assert report.content_hash in str(sheet.cell(row=2, column=1).value)


def test_the_monthly_history_marks_which_months_have_an_outcome(report):
    """A blank default rate on a matured month and a blank one on an open
    month are opposite facts, and the column is what tells them apart."""
    months = dash.available_months(APP)[-8:]
    body = report_xlsx.write(report, history_months=months)
    book = load_workbook(io.BytesIO(body))
    rows = list(book["MONTHLY HISTORY"].iter_rows(min_row=5,
                                                  values_only=True))
    seen = {row[0]: row[1] for row in rows if row[0]}
    assert set(seen) == set(months)
    for month, state in seen.items():
        expected = "MATURED" if synth.matured(month) else "NOT MATURED"
        assert str(state).startswith(expected), month


def test_an_open_month_shows_no_default_rate_at_all(report):
    """Not a zero. Not an estimate. Nothing, with the reason in the column
    beside it."""
    months = [m for m in dash.available_months(APP) if not synth.matured(m)]
    body = report_xlsx.write(report, history_months=months[:2])
    book = load_workbook(io.BytesIO(body))
    for row in book["MONTHLY HISTORY"].iter_rows(min_row=5,
                                                 values_only=True):
        if not row[0]:
            continue
        assert row[6] == "—", f"{row[0]} printed an outcome: {row[6]!r}"
        assert row[2], "the population is still reported"


def test_the_workbook_states_that_csi_cutoffs_are_a_convention(report):
    body = report_xlsx.write(report)
    book = load_workbook(io.BytesIO(body))
    note = str(book["CSI"].cell(row=2, column=1).value)
    assert "convention" in note and "not a regulatory" in note


def test_both_scorecards_produce_a_complete_report():
    """The behavioural side is a different population and a different
    equation, and has been known to be the one nobody checks."""
    for scorecard_type in (APP, BEH):
        built = rpt.build(scorecard_type, generated_by="t")
        assert rpt.coverage(built)["complete"], scorecard_type
        assert report_docx.write(built)
        assert report_xlsx.write(built)
