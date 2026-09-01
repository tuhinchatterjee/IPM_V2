"""
Report generation and download through the real routes. §51-§56, §82, §83.

Generation and download are separate routes and the tests treat them as
separate claims. Generation records what was reported and to whom;
download reproduces it. A test that only exercised download would prove
the file renders and nothing about whether anybody could later say what
was issued.
"""

from __future__ import annotations

import io

import pytest
from docx import Document
from openpyxl import load_workbook
from sqlalchemy import text as sql

from backend.scorecard import registry as reg
from backend.scorecard import report as rpt
from backend.scorecard import synthetic as synth
from tests.conftest import database_available

db = pytest.mark.skipif(not database_available(),
                        reason="needs the platform database")

APP = "APPLICATION"
DOCX = ("application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.document")
XLSX = ("application/vnd.openxmlformats-officedocument"
        ".spreadsheetml.sheet")

_OWNED = ("scorecard_report_evidence", "scorecard_reports",
          "scorecard_dashboard_pins", "scorecard_model_approvals",
          "scorecard_findings", "scorecard_validation_runs",
          "scorecard_policy_limits", "scorecard_binning_specs",
          "scorecard_model_variables", "scorecard_models")


def headers(role: str = "ADMIN", user_id: int = 1) -> dict[str, str]:
    return {"X-IPM-Role": role, "X-IPM-User-Id": str(user_id)}


@pytest.fixture(scope="module")
def client():
    from fastapi.testclient import TestClient

    from backend.api.main import app

    return TestClient(app)


@pytest.fixture
def registered():
    from backend.db.engine import SessionLocal

    session = SessionLocal()
    for table in _OWNED:
        session.execute(sql(f"DELETE FROM {table} WHERE tenant = ''"))
    session.commit()
    reg.seed(session, created_by="test")
    session.commit()
    try:
        yield session
    finally:
        session.rollback()
        for table in _OWNED:
            session.execute(sql(f"DELETE FROM {table} WHERE tenant = ''"))
        session.commit()
        session.close()


# --------------------------------------------------------------- generation


@db
def test_generating_a_report_returns_its_coverage(client, registered):
    """§89 travels with the report rather than being a separate audit."""
    body = client.post(f"/api/v1/scorecard/reports/{APP}", headers=headers(),
                       json={}).json()
    assert body["coverage"]["complete"] is True
    assert body["coverage"]["missing"] == []
    assert len(body["sections"]) >= 25
    assert body["origin"] == synth.ORIGIN


@db
def test_a_generated_report_is_recorded_with_its_disclaimer(client,
                                                            registered):
    """§0/§56. The copy somebody was handed has to be recoverable."""
    body = client.post(f"/api/v1/scorecard/reports/{APP}", headers=headers(),
                       json={}).json()
    listed = client.get(f"/api/v1/scorecard/reports/{APP}",
                        headers=headers()).json()
    assert listed["count"] == 1
    stored = listed["reports"][0]
    assert stored["report_id"] == body["report_id"]
    assert stored["opinion"] == body["opinion"]
    assert stored["generated_by"] == "ADMIN#1"

    from backend.models.platform import ScorecardReport
    row = registered.query(ScorecardReport).filter(
        ScorecardReport.report_id == body["report_id"]).one()
    assert "does not provide regulatory certification" in row.disclaimer


@db
def test_the_evidence_index_is_stored_alongside_the_report(client,
                                                           registered):
    """§55. Every figure the report printed, and the run behind it."""
    body = client.post(f"/api/v1/scorecard/reports/{APP}", headers=headers(),
                       json={}).json()
    evidence = client.get(
        f"/api/v1/scorecard/reports/evidence/{body['report_id']}",
        headers=headers()).json()
    assert evidence["count"] == body["evidence_count"] > 0
    for item in evidence["evidence"]:
        assert item["section"] and item["label"] and item["workbook_sheet"]


@db
def test_regenerating_the_same_month_keeps_the_same_content_hash(client,
                                                                  registered):
    """§82. A regeneration whose figures did not move should be visibly the
    same report, not a new unknown."""
    first = client.post(f"/api/v1/scorecard/reports/{APP}",
                        headers=headers(), json={"record": False}).json()
    second = client.post(f"/api/v1/scorecard/reports/{APP}",
                         headers=headers(role="DATA_STEWARD", user_id=2),
                         json={"record": False}).json()
    assert first["content_hash"] == second["content_hash"]


def test_only_a_report_role_may_generate_one(client):
    """§87. An Analyst reads the dashboards; issuing a validation report is
    a narrower act."""
    for role in ("ANALYST", "VIEWER"):
        response = client.post(f"/api/v1/scorecard/reports/{APP}",
                               headers=headers(role=role), json={})
        assert response.status_code == 403, role


def test_an_unknown_scorecard_is_refused(client):
    response = client.post("/api/v1/scorecard/reports/MORTGAGE",
                           headers=headers(), json={})
    assert response.status_code in (404, 422)


def test_a_report_can_be_produced_without_a_registry(client):
    """A workspace that built a lake and registered nothing can still
    validate it — and section 1 says the owner is unavailable rather than
    inventing one."""
    from backend.db.engine import SessionLocal

    session = SessionLocal()
    for table in _OWNED:
        session.execute(sql(f"DELETE FROM {table} WHERE tenant = ''"))
    session.commit()
    session.close()

    body = client.post(f"/api/v1/scorecard/reports/{APP}", headers=headers(),
                       json={}).json()
    assert body["coverage"]["complete"] is True
    assert body["model_version"] == "unregistered"
    cover = next(s for s in body["sections"] if s["number"] == "1")
    assert "not recorded in the scorecard model registry" in cover["narrative"]


# ----------------------------------------------------------------- download


def test_the_word_download_is_a_real_docx_with_the_named_filename(client):
    response = client.get(
        f"/api/v1/scorecard/reports/{APP}/download?fmt=docx",
        headers=headers())
    assert response.status_code == 200
    assert response.headers["content-type"] == DOCX
    disposition = response.headers["content-disposition"]
    assert "CreditProbe_APPLICATION_INCUMBENT_" in disposition
    assert disposition.endswith('_Validation_Report.docx"')
    assert response.headers["x-creditprobe-origin"] == synth.ORIGIN
    assert response.headers["x-creditprobe-content-hash"]

    document = Document(io.BytesIO(response.content))
    headings = [p.text for p in document.paragraphs
                if p.style.name.startswith("Heading")]
    assert "1 Cover and document control" in headings
    assert "12 Overall validation conclusion" in headings


def test_the_evidence_download_is_a_real_workbook(client):
    response = client.get(
        f"/api/v1/scorecard/reports/{APP}/download?fmt=xlsx&history_months=6",
        headers=headers())
    assert response.status_code == 200
    assert response.headers["content-type"] == XLSX
    book = load_workbook(io.BytesIO(response.content))
    assert book.sheetnames[0] == "EVIDENCE INDEX"
    assert "REGULATORY MAPPING" in book.sheetnames


def test_an_unknown_format_is_refused_rather_than_defaulted(client):
    """Quietly serving a Word file to somebody who asked for a PDF is how a
    download button ends up lying about what it produces."""
    response = client.get(
        f"/api/v1/scorecard/reports/{APP}/download?fmt=pdf",
        headers=headers())
    assert response.status_code == 400


def test_a_month_with_an_open_window_downloads_and_says_so(client):
    """§7 through the whole stack: build, render, download."""
    from backend.scorecard import dashboard as dash

    open_months = [m for m in dash.available_months(APP)
                   if not synth.matured(m)]
    assert open_months, "no month with an open performance window"

    response = client.get(
        f"/api/v1/scorecard/reports/{APP}/download?fmt=docx"
        f"&month={open_months[0]}", headers=headers())
    assert response.status_code == 200
    prose = " ".join(p.text for p in
                     Document(io.BytesIO(response.content)).paragraphs)
    assert "Not reported." in prose
    assert "performance window" in prose
    assert synth.window_closes(open_months[0]) in prose


def test_the_report_and_the_download_agree(client):
    """A download that disagrees with the screen it was started from is a
    worse failure than a slow one."""
    body = client.post(f"/api/v1/scorecard/reports/{APP}", headers=headers(),
                       json={"record": False}).json()
    response = client.get(
        f"/api/v1/scorecard/reports/{APP}/download?fmt=docx",
        headers=headers())
    assert response.headers["x-creditprobe-content-hash"] == \
        body["content_hash"]


def test_the_filename_helper_and_the_route_agree():
    built = rpt.build(APP, generated_by="t")
    assert rpt.filename_for(built, "xlsx").endswith(
        "_Validation_Report.xlsx")
