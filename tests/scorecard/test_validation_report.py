"""The validation report, and the ways a report lies.

A report is where every earlier care can be undone in one paragraph. The
tests here are about three of those paragraphs: an opinion that outruns its
evidence, a refusal that disappears between the engine and the page, and a
number in a document that no longer matches the run it came from.
"""

from __future__ import annotations

import io

import pytest

from backend.scorecard.validation import (
    findings,
    models,
    registry,
    runner,
    states,
)
from backend.scorecard.validation import (
    report as scv_report,
)

STAMP = "2026-09-05T00:00:00+00:00"


@pytest.fixture(scope="module")
def champion() -> models.Model:
    return models.get("sme_champion")


@pytest.fixture(scope="module")
def results(champion: models.Model) -> list[states.Result]:
    out: list[states.Result] = []
    for category in registry.CATEGORIES:
        out.extend(runner.run_category(category, champion))
    return out


@pytest.fixture(scope="module")
def built(champion: models.Model, results: list[states.Result]):
    return scv_report.build(champion, results, generated_at=STAMP)


def _pass(test_id: str) -> states.Result:
    return states.measured(
        test_id, states.PASS, 0.9, limit=0.5, limit_source="TEST",
        detail="inside its limit", observations=10_000, events=800)


# ------------------------------------------- the opinion cannot outrun itself


def test_a_report_that_measured_almost_nothing_has_no_opinion(
        champion: models.Model) -> None:
    """USE AS IS on four tests of forty-eight is true and misleading."""
    thin = [_pass("DISC-AUC")] + [
        states.not_matured(t.test_id, period="2026-01", closes="2027-01")
        for t in registry.TESTS if t.test_id != "DISC-AUC"]
    made = scv_report.build(champion, thin, generated_at=STAMP)
    assert made.opinion == scv_report.INSUFFICIENT_EVIDENCE
    assert "%" in made.section("2").narrative


def test_a_clean_run_reads_use_as_is(champion: models.Model) -> None:
    clean = [_pass(t.test_id) for t in registry.TESTS]
    made = scv_report.build(champion, clean, generated_at=STAMP)
    assert made.opinion == scv_report.USE_AS_IS


def test_a_failed_replication_stops_the_report_relying_on_itself(
        champion: models.Model) -> None:
    """Every other result describes a model that is not in production."""
    broken = [_pass(t.test_id) for t in registry.TESTS
              if t.test_id != "IMPL-REPLICATE"]
    broken.append(states.measured(
        "IMPL-REPLICATE", states.FAIL, 0.4, limit=0.0,
        limit_source="STRUCTURAL",
        detail="two in five rows do not reproduce", observations=10_000))
    made = scv_report.build(champion, broken, generated_at=STAMP)
    assert made.opinion == scv_report.DO_NOT_USE_UNTIL_REMEDIATED


def test_findings_make_the_opinion_conditional(built) -> None:
    assert built.opinion == scv_report.USE_WITH_CONDITIONS
    assert "finding" in built.section("2").narrative


def test_no_opinion_claims_approval() -> None:
    """Approval is a committee's act, not a document's."""
    for opinion in scv_report.OPINIONS:
        assert "APPROV" not in opinion.upper()
    joined = " ".join(scv_report.OPINION_MEANING.values()).upper()
    assert "IS APPROVED" not in joined


# ------------------------------------------------- the refusals reach the page


def test_every_result_appears_in_the_document(
        built, results: list[states.Result]) -> None:
    """A report that omits what it could not measure hides its own scope."""
    printed = {row[0] for section in built.sections
               for table in section.tables for row in table.rows}
    for result in results:
        assert result.test_id in printed, (
            f"{result.test_id} ran and does not appear in the report")


def test_an_unmeasured_row_shows_a_dash_not_a_zero(built) -> None:
    """A zero reads as a measurement and a blank reads as unfinished."""
    labels = {states.STATE_LABELS[s] for s in states.UNMEASURED}
    for section in built.sections:
        for table in section.tables:
            if table.columns != scv_report.RESULT_COLUMNS:
                continue
            for row in table.rows:
                if row[4] in labels:
                    assert row[2] == "—", (
                        f"{row[0]} is {row[4]} and shows {row[2]!r}")


def test_a_row_with_no_limit_says_so_rather_than_showing_a_number(
        built) -> None:
    for section in built.sections:
        for table in section.tables:
            if table.columns != scv_report.RESULT_COLUMNS:
                continue
            for row in table.rows:
                assert row[3] != "None"
                assert row[3] != ""


def test_the_limitations_section_comes_before_the_results(built) -> None:
    """A limitation at the back of a document has been read past."""
    numbers = [s.number for s in built.sections]
    assert numbers.index("4") < numbers.index("6")


# ------------------------------------------------- the document is the run


def test_the_same_results_build_the_same_report(
        champion: models.Model, results: list[states.Result]) -> None:
    first = scv_report.build(champion, results, generated_at=STAMP)
    second = scv_report.build(champion, results, generated_at=STAMP)
    assert first.content_hash == second.content_hash


def test_a_changed_result_changes_the_hash(
        champion: models.Model, results: list[states.Result]) -> None:
    """Otherwise the hash answers 'was this run twice?' and nothing else."""
    before = scv_report.build(champion, results, generated_at=STAMP)
    after = scv_report.build(
        champion, [*results[:-1], _pass("DISC-AUC")], generated_at=STAMP)
    assert before.content_hash != after.content_hash


def test_a_new_generation_stamp_does_not_change_the_hash(
        champion: models.Model, results: list[states.Result]) -> None:
    """The hash is about the assessment, not the act of generating it."""
    first = scv_report.build(champion, results, generated_at=STAMP)
    later = scv_report.build(champion, results,
                             generated_at="2027-01-01T00:00:00+00:00")
    assert first.content_hash == later.content_hash


def test_every_measured_figure_is_in_the_evidence_register(
        built, results: list[states.Result]) -> None:
    measured = {r.test_id for r in results if r.measured}
    registered = {e.metric for e in built.evidence}
    assert measured == registered


def test_the_evidence_register_carries_no_refusals(built) -> None:
    for entry in built.evidence:
        assert entry.validation_state in states.MEASURED
        assert entry.value is not None
        assert entry.method, f"{entry.metric} cites no method"


def test_the_report_id_is_a_window_not_a_concatenation(built) -> None:
    """The periods on results are a mix of ranges and single months.

    Sorting those as strings and joining the ends produced a report id
    reading `2023-01..2024-04..2025-12`, which is not a window. Both windows
    now come from the model's own data.
    """
    assert built.report_id.count("..") <= 1
    assert built.period.count("..") <= 1


def test_the_cover_states_both_windows(built) -> None:
    """Outcome tests and stability tests run on different periods."""
    control = {row[0] for row in built.sections[0].tables[0].rows}
    assert any("Matured window" in item for item in control)
    assert any("Latest data period" in item for item in control)


# -------------------------------------------------------------- the document


def test_the_docx_opens_and_carries_the_report(built) -> None:
    """Not 'a file was produced'. Opened, and read back."""
    import docx

    blob = scv_report.docx(built)
    assert len(blob) > 20_000
    document = docx.Document(io.BytesIO(blob))

    headings = [p.text for p in document.paragraphs
                if p.style.name.startswith("Heading")]
    for section in built.sections:
        if not section.has_content:
            continue
        assert any(section.title in h for h in headings), (
            f"section {section.number} {section.title!r} is not in the .docx")
    assert len(document.tables) >= len(
        [t for s in built.sections for t in s.tables])


def test_the_docx_contains_no_orphan_zero(built) -> None:
    """The failure this whole engine exists to prevent, checked on the page."""
    import docx

    document = docx.Document(io.BytesIO(scv_report.docx(built)))
    labels = {states.STATE_LABELS[s] for s in states.UNMEASURED}
    for table in document.tables:
        header = [c.text for c in table.rows[0].cells]
        if header != scv_report.RESULT_COLUMNS:
            continue
        for row in table.rows[1:]:
            cells = [c.text for c in row.cells]
            if cells[4] in labels:
                assert cells[2] == "—", (
                    f"{cells[0]} reached the page as {cells[2]!r} while "
                    f"being {cells[4]}")


def test_the_findings_reach_the_document_with_their_remediation(
        built, champion: models.Model,
        results: list[states.Result]) -> None:
    assessed = findings.assess(results, champion)
    printed = {row[0] for section in built.sections
               for table in section.tables for row in table.rows}
    for made in assessed:
        assert made.finding_id in printed
    section = built.section("5")
    for row in section.tables[0].rows:
        assert row[4].strip(), f"{row[0]} reaches the report with no remedy"
        assert row[5].strip(), f"{row[0]} says nothing about verification"


def test_the_regulatory_section_carries_its_disclaimer(built) -> None:
    from backend.scorecard.validation import regulatory

    assert built.section("7").narrative == regulatory.DISCLAIMER


def test_the_document_says_draft_where_a_reader_will_see_it(built) -> None:
    """Not only on the button that produced it.

    A generated DOCX was inspected during this build and the word "draft"
    appeared nowhere in it — the screen said draft, the file did not. A
    document that does not announce itself as a draft is the exact artefact
    that ends up in a committee pack with somebody's name under it, and by
    then the screen it came from is long gone.
    """
    import docx

    assert "DRAFT" in built.title

    document = docx.Document(io.BytesIO(scv_report.docx(built)))
    everywhere = "\n".join(p.text for p in document.paragraphs) + "\n".join(
        cell.text for table in document.tables
        for row in table.rows for cell in row.cells)
    assert everywhere.count("DRAFT") >= 2, (
        "the cover and the document-control table both say it")
    assert "does not issue validation opinions" in everywhere
    # The one word this document must never contain.
    assert "compliant" not in everywhere.lower()
